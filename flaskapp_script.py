from flask import Flask, request, jsonify, send_from_directory
from docx import Document
from datetime import datetime
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from werkzeug.utils import secure_filename
import os
from docx2pdf import convert
from jira import JIRA
import requests
from requests.auth import HTTPBasicAuth
import math 
import pythoncom
import tempfile
import shutil
import io
from docx2pdf import convert as docx_to_pdf_convert
from config import JIRA_API_TOKEN, CONFLUENCE_API_TOKEN, SP_CLIENT_ID, SP_CLIENT_SECRET, SP_TENANT_ID

SHAREPOINT_SITE_ID = 'your sharepoint site id'
DOCUMENT_PATH = 'document path'

app = Flask(__name__)

@app.route('/')
def api_root():
    return 'Welcome to the JIRA Webhook Receiver'
    
@app.route('/jira', methods=['POST'])
def api_jira_message():
    if request.headers.get('Content-Type', '').lower() == 'application/json':
        data = request.json.get('issue', {})
        items = extract_items(data)
        
        # Fetch the Word template from Confluence
        template_content = fetch_confluence_template()
        
        if template_content:
            access_token = get_access_token()  # Make sure this token is correctly retrieved
            if not access_token:
                return jsonify({"error": "Failed to retrieve access token"}), 500
            
            upload_result = create_word_document(data, items, template_content, access_token)
            
            if upload_result == "Upload successful":
                return jsonify({"message": "Documents processed and uploaded successfully"}), 200
            else:
                return jsonify({"error": upload_result}), 500
        else:
            return jsonify({"error": "Failed to fetch template from Confluence"}), 500
    else:
        return jsonify({"error": "Invalid Content-Type, expected 'application/json'"}), 400


def fetch_confluence_template():
    base_url = 'https://yoursite.atlassian.net/wiki'
    page_id = '3481468945'  # Ensure this is your correct page ID
    headers = {
        'Authorization': 'Bearer ' + CONFLUENCE_API_TOKEN,
        'Accept': 'application/json'
    }
    api_url = f'{base_url}/rest/api/content/{page_id}/child/attachment'
    response = requests.get(api_url, headers=headers)
    if response.status_code == 200:
        attachments = response.json().get('results', [])
        for attachment in attachments:
            if attachment['title'] == "New Quote Template.docx":
                download_url = base_url + attachment['_links']['download']
                response = requests.get(download_url, headers=headers)
                if response.status_code == 200:
                    return io.BytesIO(response.content)
                else:
                    print("Failed to download template:", response.status_code, response.text)
                    return None
        print("Template 'New Quote Template.docx' not found among attachments.")
        return None
    else:
        print(f"Failed to fetch attachments: {response.status_code} {response.text}")
        return None
    
def get_access_token():
    tenant_id = SP_TENANT_ID
    client_id = SP_CLIENT_ID
    client_secret = SP_CLIENT_SECRET
    scope = 'https://graph.microsoft.com/.default'

    url = f'https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token'
    payload = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': scope
    }

    response = requests.post(url, data=payload)
    response_data = response.json()
    if response.status_code == 200:
        return response_data.get('access_token')
    else:
        print("Failed to retrieve access token:", response_data.get('error_description'))
        return None

    
def upload_to_sharepoint(file_stream, filename, folder_name, access_token):
    site_id = '2eb916ed-b02a-40b4-965f-38831bea5688'  # Ensure this is your correct SharePoint site ID
    upload_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{folder_name}/{filename}:/content"
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/octet-stream'
    }
    file_stream.seek(0)  # Ensure the stream is at the beginning
    response = requests.put(upload_url, headers=headers, data=file_stream)
    if response.status_code == 201:
        print("File uploaded successfully to SharePoint.")
        return "Upload successful"
    else:
        print(f"Failed to upload file to SharePoint: {response.status_code} {response.text}")
        return "Upload failed"
    
def extract_items(data):
    items = []
    for i in range(1, 6):  # Assuming there are 5 items at most
        item_key = f'item{i}'
        if item_key in data and data[item_key]:
            item_number = data.get(item_key)
            description_key = f'itemDescrip{i}'
            description = data.get(description_key, '')
            qty_str = data.get(f'qty{i}', '0').strip()
            qty = float(qty_str) if qty_str else 0
            unit_key = f'Unit_{i}'
            unit = data.get(unit_key, 'EA')
            price_str = data.get(f'price{i}', '0.0').strip()
            price = float(price_str) if price_str else 0.0
            total = qty * price
            items.append({
                "Item#": item_number,
                "Description": description,
                "Qty": qty,
                "Unit": unit,
                "Unit Price": f"${price:,.2f}",
                "Total": f"${total:,.2f}"
            })
    return items

def create_word_document(data, items, template_content, access_token):
    try:
        today_date = datetime.now().strftime("%Y-%m-%d")
        client_name = data.get('clientName', 'Unknown_Client')
        issue_key = data.get('key', 'Unknown_Key')
        filename = f"{client_name}_{issue_key}"

        # Load the template document from bytes
        document = Document(template_content)
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        populate_table_header(table.rows[0])
        total_sum = populate_table_data(table, items)
        phi_collect(table, items)
        phi_shipping(table, items, data)
        add_final_row(table, total_sum)
        replace_placeholders(document, data, today_date)

        # Replace the {{items}} placeholder and insert the table
        found_placeholder = False
        for paragraph in document.paragraphs:
            if '{{items}}' in paragraph.text:
                paragraph.text = ''
                p = paragraph._p  # Get the XML element for the paragraph
                p.addnext(table._tbl)  # Insert the table XML element after the paragraph XML element
                found_placeholder = True
                break
        if not found_placeholder:
            print("Warning: '{{items}}' placeholder not found in the document. Table appended at the end of the document.")
            document.add_page_break()
            new_paragraph = document.add_paragraph()
            new_paragraph._p.addnext(table._tbl)

        # Save the document to a Word file
        word_filename = f"{filename}.docx"
        document.save(word_filename)

        # Convert the Word document to PDF
        pdf_filename = f"{filename}.pdf"
        convert_to_pdf(word_filename, pdf_filename)

        # Upload both Word and PDF documents to SharePoint
        with open(word_filename, 'rb') as word_file, open(pdf_filename, 'rb') as pdf_file:
            word_stream = io.BytesIO(word_file.read())
            pdf_stream = io.BytesIO(pdf_file.read())
            upload_result_word = upload_to_sharepoint(word_stream, word_filename, 'Quotes', access_token)
            upload_result_pdf = upload_to_sharepoint(pdf_stream, pdf_filename, 'Quotes', access_token)
            
            # Attach the PDF AND Word Doc to the JIRA issue
            if upload_result_pdf == "Upload successful":
                attach_pdf_to_jira_issue(issue_key, pdf_filename)
            if upload_result_word == "Upload successful":
                attach_word_doc_jira(issue_key, word_filename)

        # Cleanup: delete temporary files
        os.remove(word_filename)
        os.remove(pdf_filename)
    
        # Check if the total sum is greater than $4000 and then call approved_quote if needed
        if total_sum > 4000:
            print(f"Total sum is {total_sum}, which is greater than $4000. Running approved_quote.")
            approved_quote(issue_key)
        else:
            print(f"Total sum is {total_sum}, which is not greater than $4000. Skipping approved_quote.")
            needs_review(issue_key, total_sum, connect_to_jira())

        return "Upload successful"
    except Exception as e:
        print(f"Error in create_word_document: {e}")
        return "Error creating document"

def populate_table_header(row):
    headers = ["Item#", "Description", "Qty", "Unit", "Unit Price", "Total"]
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        cell.text = header
        set_cell_background(cell, "2887dd")  # Set blue background
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Cm(4) if header != "Description" else Cm(10)

def populate_table_data(table, items):
    total_sum = 0
    for item in items:
        row = table.add_row().cells
        for idx, key in enumerate(["Item#", "Description", "Qty", "Unit", "Unit Price", "Total"]):
            if key == "Qty":  # Check if the current cell is for the Qty column
                row[idx].text = str(int(item[key]))  # Convert quantity to integer before displaying
            else:
                row[idx].text = str(item[key])
            paragraph = row[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if key == "Total":
                total_sum += float(item[key].strip('$').replace(',', ''))
    return total_sum

def add_final_row(table, total_sum):
    row = table.add_row().cells
    row[4].text = "Total"
    row[5].text = f"${total_sum:,.2f}"

    for cell in row:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        remove_borders(cell)

def set_cell_background(cell, fill):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading)

#Attempting to Remove Inside Borders from the last row for the ALL_Total
def remove_borders(cell):
    for border in ["top", "left", "bottom", "right"]:
        tag = 'w:' + border + 'Border'
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'nil')
        cell._element.get_or_add_tcPr().append(element)
        
# IF PhiCollect equals "YES"

def phi_collect(table, items):
    phi_collect_value = "yes"  
    if phi_collect_value.lower() == "yes":
        row = table.add_row().cells
        for i in range(6):
            row[i].text = ""
        row[0].text = "WMS-800"
        row[1].text = "Water Sample Collection"
        qty_sum = sum(item["Qty"] for item in items)
        if qty_sum < 20:
            row[2].text = "1"
        else:
            row[2].text = str(int(qty_sum))
        row[3].text = "EA"
        if row[2].text == "1":
            row[4].text = "$600.00"
        else:
            row[4].text = "$30.00"
        total_price = float(row[4].text.strip('$')) * int(row[2].text)
        row[5].text = f"${total_price:,.2f}"
        for cell in row:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
#IF PhiShipping equals "YES"

def phi_shipping(table, items, data):
    phi_shipping_value = "yes"  # Or from some configuration or data field
    if phi_shipping_value.lower() == "yes":
        row = table.add_row().cells
        for i in range(6):
            row[i].text = ""
        row[0].text = "WMS-9970"
        row[1].text = "Shipping and Handling - Overnight Return"
        total_divided_qty = sum(float(item['Qty']) / float(data.get(f'itemMAX_{i+1}', 1)) for i, item in enumerate(items))
        total_divided_qty_rounded = math.ceil(total_divided_qty)  # Use math.ceil to round up
        row[2].text = str(total_divided_qty_rounded)
        row[3].text = "BOX"
        row[4].text = "$110.00"
        total_price = 110.00 * total_divided_qty_rounded
        row[5].text = f"${total_price:,.2f}"
        for cell in row:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def replace_placeholders(document, data, today_date):
    for paragraph in document.paragraphs:
        if '{{clientName}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{clientName}}', data.get('clientName', 'Unknown Client'))
        if '{{pocName}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{pocName}}', data.get('pocName', ''))
        if '{{title}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{title}}', data.get('title', ''))
        if '{{clientCode}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{clientCode}}', data.get('clientCode', ''))
        if '{{today_date}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{today_date}}', today_date)
        if '{{issue_Key}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{issue_Key}}', data.get('key', ''))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if '{{clientName}}' in cell_text:
                    cell.text = cell_text.replace('{{clientName}}', data.get('clientName', 'Unknown Client'))
                if '{{pocName}}' in cell_text:
                    cell.text = cell_text.replace('{{pocName}}', data.get('pocName', ''))
                if '{{title}}' in cell_text:
                    cell.text = cell_text.replace('{{title}}', data.get('title', ''))
                if '{{clientCode}}' in cell_text:
                    cell.text = cell_text.replace('{{clientCode}}', data.get('clientCode', ''))
                if '{{today_date}}' in cell_text:
                    cell.text = cell_text.replace('{{today_date}}', today_date)
                if '{{shippingAddress}}' in cell_text:
                    cell.text = cell_text.replace('{{shippingAddress}}', data.get('shippingAddress', ''))
                if '{{address}}' in cell_text:
                    cell.text = cell_text.replace('{{address}}', data.get('address', ''))

def replace_header(document, data):
    for section in document.sections:
        for header in section.header_parts:
            for paragraph in header.paragraphs:
                if '{{issue_Key}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{issue_Key}}', data.get('key', ''))
                    
## Function to connect to JIRA ##
def connect_to_jira():
    # Directly define credentials
    username = 'user@email.com'
    api_token = JIRA_API_TOKEN
    
    # Set up basic authentication
    auth = HTTPBasicAuth(username, api_token)
    return auth

## Function to add attachments to a JIRA issue ##
def attach_pdf_to_jira_issue(issue_key, pdf_filename):
    try:
        # Set up the URL for the JIRA REST API
        url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/attachments"

        # Connect to JIRA
        auth = connect_to_jira()

        # Set up the headers
        headers = {
            "X-Atlassian-Token": "no-check",
            "Accept": "application/json"
        }

        # Open the PDF file in binary mode for uploading
        with open(pdf_filename, 'rb') as pdf_file:
            files = {
                "file": (os.path.basename(pdf_filename), pdf_file, 'application/pdf')
            }

            # Make the request to add the attachment
            response = requests.post(url, headers=headers, files=files, auth=auth)

            # Print the response from JIRA
            if response.ok:
                print("Attachment uploaded successfully!")
                print(response.json())  # This will print the response JSON from JIRA
            else:
                print(f"Failed to upload attachment: {response.status_code} {response.text}")

    except Exception as e:
        print(f"Error uploading attachment to JIRA: {e}")
        raise

def attach_word_doc_jira(issue_key, word_filename):
    try:
        # Set up the URL for the JIRA REST API
        url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/attachments"

        # Connect to JIRA
        auth = connect_to_jira()

        # Set up the headers
        headers = {
            "X-Atlassian-Token": "no-check",
            "Accept": "application/json"
        }

        # Open the Word file in binary mode for uploading
        with open(word_filename, 'rb') as word_file:
            files = {
                "file": (os.path.basename(word_filename), word_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }

            # Make the request to add the attachment
            response = requests.post(url, headers=headers, files=files, auth=auth)

            # Print the response from JIRA
            if response.ok:
                print("Attachment uploaded successfully!")
                print(response.json())  # This will print the response JSON from JIRA
            else:
                print(f"Failed to upload attachment: {response.status_code} {response.text}")

    except Exception as e:
        print(f"Error uploading attachment to JIRA: {e}")
        raise

def convert_to_pdf(docx_path, pdf_path):
    """
    Convert a DOCX file to a PDF file using python-docx and convert docx.
    """
    try:
        # Initialize the COM library for converting DOCX to PDF
        pythoncom.CoInitialize()

        # Convert DOCX to PDF
        docx_to_pdf_convert(docx_path, pdf_path)

        # Check if the PDF was created successfully
        if os.path.exists(pdf_path):
            # Uninitialize the COM library
            pythoncom.CoUninitialize()
            return True
        else:
            print("Failed to convert DOCX to PDF.")
            # Uninitialize the COM library
            pythoncom.CoUninitialize()
            return False
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        # Uninitialize the COM library
        pythoncom.CoUninitialize()
        return False
    
    
## Function designed to post downloaded hyperlinks for both PDF and Word Doc into a comment ##
    # only runs if Total_Sum is GREATER than $4000 #
    
def approved_quote(issue_key):
    try:
        url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}?fields=attachment"
        auth = connect_to_jira()
        headers = {"Accept": "application/json"}
        response = requests.get(url, headers=headers, auth=auth)
        
        if response.ok:
            data = response.json()
            attachments = data['fields']['attachment']
            if attachments:
                post_comment_to_jira(issue_key, attachments)
                transition_issue_to_completed(issue_key)
            else:
                print("No attachments found for the issue.")
        else:
            print(f"Failed to get attachments: {response.status_code} {response.text}")
    except Exception as e:
        print(f"Error in approved_quote: {e}")

def post_comment_to_jira(issue_key, attachments):
    try:
        url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/comment"
        auth = connect_to_jira()
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }

        # Prepare the body of the comment
        body = {
            "body": {
                "version": 1,
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [
                            {
                                "type": "text",
                                "text": "Your Quote has been successfully generated:"
                            }
                        ]
                    }
                ]
            }
        }

        # Append attachments to the comment
        for attachment in attachments:
            file_id = attachment['id']  # Correctly access the 'id' field of each attachment
            file_name = attachment['filename']
            media_item = {
                "type": "mediaSingle",
                "attrs": {"layout": "center"},
                "content": [{
                    "type": "media",
                    "attrs": {
                        "type": "file",
                        "id": file_id,
                        "collection": issue_key
                    }
                }]
            }
            body['body']['content'].append({
                "type": "paragraph",
                "content": [{
                    "type": "text",
                    "text": f"{file_name} - {file_id}",
                    "marks": [{
                        "type": "link",
                        "attrs": {
                            "href": f"https://yoursite.atlassian.net/secure/attachment/{file_id}/{file_name}"
                        }
                    }]
                }]
            })

        # Send the comment post request
        response = requests.post(url, json=body, headers=headers, auth=auth)
        if response.status_code in [200, 201]:
            print("Comment with attachments posted successfully!")
        else:
            print(f"Failed to post comment: {response.status_code} - {response.text}")

    except Exception as e:
        print(f"Error posting comment to JIRA: {e}")
        
## if approved quote runs then ticket status should transition to 'Completed' ##
        
def get_completed_transition_id(issue_key):
    url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/transitions"
    auth = connect_to_jira()
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers, auth=auth)

    if response.ok:
        transitions = response.json().get('transitions', [])
        for transition in transitions:
            if transition['name'].lower() == 'completed':
                return transition['id']
    return None

def transition_issue_to_completed(issue_key):
    transition_id = get_completed_transition_id(issue_key)
    if transition_id:
        url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/transitions"
        auth = connect_to_jira()
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        payload = {
            "transition": {
                "id": transition_id
            }
        }
        response = requests.post(url, json=payload, headers=headers, auth=auth)
        if response.status_code in [200, 204]:
            print("Issue transitioned to Completed successfully!")
        else:
            print(f"Failed to transition issue: {response.status_code} - {response.text}")
    else:
        print("Transition ID for 'Completed' could not be found.")
        

## IF total_sum is LESS than $4000 ##
def needs_review(issue_key, total_sum, auth):
    if total_sum < 4000:
        # Post a comment to the JIRA issue
        post_comment(issue_key, "Due to policy, quote needs to be reviewed", auth)
        # Assign the ticket to Nancy Galvez using her email
        assign_ticket(issue_key, "ngalvez@yoursite.com", auth)

def post_comment(issue_key, message, auth):
    url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/comment"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }
    payload = {
        "body": {
            "version": 1,
            "type": "doc",
            "content": [{
                "type": "paragraph",
                "content": [{
                    "type": "text",
                    "text": message
                }]
            }]
        }
    }
    response = requests.post(url, json=payload, headers=headers, auth=auth)
    if response.status_code in [200, 201]:
        print("Comment posted successfully!")
    else:
        print(f"Failed to post comment: {response.status_code} - {response.text}")

def get_account_id_by_email(email, auth):
    # Fetch accountId based on user's email
    url = f"https://yoursite.atlassian.net/rest/api/3/user/search?query={email}"
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers, auth=auth)
    if response.status_code == 200 and response.json():
        return response.json()[0]['accountId']  # Assuming the first user returned is the correct one
    else:
        print(f"Failed to get accountId: {response.status_code} - {response.text}")
        return None

def assign_ticket(issue_key, assignee_email, auth):
    url = f"https://yoursite.atlassian.net/rest/api/3/issue/{issue_key}/assignee"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }
    payload = {
        "accountId": get_account_id_by_email(assignee_email, auth)  # Fetch accountId based on user's email
    }
    response = requests.put(url, json=payload, headers=headers, auth=auth)
    if response.status_code in [200, 204]:
        print("Issue assigned successfully!")
    else:
        print(f"Failed to assign issue: {response.status_code} - {response.text}")



